# Main File (work in progress) to orchestrate the entire rag pipeline from here..
# Loading all the libraries in here
import os
from dotenv import load_dotenv
load_dotenv()
from utils.logger import setup_logger
logger=setup_logger(__name__)
import numpy as np
from utils.constants import save_documents

# Loading the documents in here..
from loader.loader import text_loader,table_loader,image_loader
print("Main.py Loading the datasets")
logger.info("Main.py Loading the datasets")
text_doc=text_loader()
print("Loaded Text Data")
table_doc=table_loader()
print("Loaded Table Data")
images_doc=image_loader()
print("Loaded images Data")
#-----------------Saving the loaded documents-------------------
data_folder="artifacts/loaded_data"
filename=data_folder+"/text_docs.pkl"
save_documents(text_doc,filename,data_folder)
filename=data_folder+"/table_docs.pkl"
save_documents(table_doc,filename,data_folder)
filename=data_folder+"/image_docs.pkl"
save_documents(images_doc,filename,data_folder)
print("Successfully loaded and saved the data")
logger.info("Successfully loaded and saved the data")
#-----------------------------------------------------------------
# 2.Chunking the documents
from Chunks.chunker import read_files,chunk_files
print("Feeding the data into the chunker")
logger.info("Feeding the data into the chunker")

text_data,table_data,images_data=read_files()

print("Chunking the data from main.py")
logger.info("Chunking the data from main.py")

text_chunks,table_chunks,images_chunks=chunk_files(text_data,table_data,images_data)
#-----------------Saving the loaded documents-------------------
data_folder="artifacts/chunked_data"
filename=data_folder+"/text_data_chunk.pkl"
save_documents(text_chunks,filename,data_folder)
filename=data_folder+"/table_data_chunk.pkl"
save_documents(table_chunks,filename,data_folder)
filename=data_folder+"/images_data_chunk.pkl"
save_documents(images_chunks,filename,data_folder)

print("Successfully loaded and saved the chunked data")
logger.info("Successfully loaded and saved the chunked data")
#-----------------------------------------------------------------
# 3. For embedding go run - `python -m Embeddings.embedding`

#-----------------------------------------------------------------
#4. Creating VectorStore (FAISS) and benchmarking the types of indices and similarity_searches
from Vector_DB.indexer import build_faiss_indices,save_indices

# Building the faiss indices using the existing chunked data
from langchain_openai import OpenAIEmbeddings
embedding_model=OpenAIEmbeddings(
    model="text-embedding-3-small"
)
# COmbining all the chunked data
all_docs=text_chunks+table_chunks+images_chunks

# Retrieving the 3 types of indices
print("Starting to create the 3 types of indices..")
logger.info("Starting to create the 3 types of indices..")
flat_index,hnsw_index,ivf_index=build_faiss_indices(all_docs,embedding_model)
print("Returned to main function")
logger.info("Returned to main function")
#-----------------Saving the embedded documents-------------------
save_indices(flat_index,hnsw_index,ivf_index)
logger.info("Saved the indices data successfully..")
print("Saved the indices data successfully..")
#-----------------------------------------------------------------
# 5. Integrating retriever module and it's metrics data
from Retrievers.retriever_pipeline import retrieve_and_compare,load_indices
flat_index,hnsw_index,ivf_index=load_indices()
# Enter the queries
n = int(input("Enter the number of queries: "))
k=int(input("Enter the number of documents you would like to retrieve: "))
test_queries = []
for i in range(n):
    query = input("Enter your query: ")
    test_queries.append(query)
#-----------------------------------------------------------------
# Step-6:-->Prompt Template
from utils.constants import TEMPLATE
from langchain.prompts import PromptTemplate

rag_prompt=PromptTemplate(
    input_variables=["context","question"],
    template=TEMPLATE
)
print("Prompt taken successfully")
logger.info("Prompt taken successfully")
#-----------------------------------------------------------------
# Step 7--> Creating a LLM Model Response
from langchain_openai import ChatOpenAI
llm=ChatOpenAI(model="gpt-4o-mini",temperature=0)
print("LLM Created successfully")
logger.info("LLM Created successfully")
#-----------------------------------------------------------------
# Step:8-->Saving into docx
from docx import Document
import os
from utils.logger import setup_logger

logger = setup_logger(__name__)

def save_to_docx(query, context, answer, k=None, fastest_index=None, filename="artifacts/output.docx"):
    # Ensure the folder exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)

    # Open existing DOCX if it exists, otherwise create new
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
        doc.add_heading("RAG Query Results", 0)

    # Add new query result
    doc.add_heading("User Query:", level=1)
    doc.add_paragraph(query)

    if k is not None:
        doc.add_heading("Number of Documents Retrieved:", level=2)
        doc.add_paragraph(f"{k}")

    if fastest_index is not None:
        doc.add_heading("Fastest Index Search:", level=2)
        doc.add_paragraph(f"{fastest_index}")

    doc.add_heading("Retrieved Context:", level=1)
    doc.add_paragraph(context[:1000] + "..." if len(context) > 1000 else context)

    doc.add_heading("LLM Answer:", level=1)
    doc.add_paragraph(answer)

    doc.add_page_break()  # separate queries

    doc.save(filename)
    print(f"Saved results to {filename}")
    logger.info(f"Saved results to {filename}")

#-----------------------------------------------------------------
# STEP 9: (FINAL) making the full pipeline
for query in test_queries:
    results = retrieve_and_compare(query, flat_index, hnsw_index, ivf_index, k=k)

    if results:
        # Find fastest index (based on retrieval time only)
        fastest_index = min(results, key=lambda x: results[x]['time_ms'])
        content_output=f"\n FASTEST: {fastest_index.upper()} with {results[fastest_index]['time_ms']:.2f} ms" #-->to be added in docx
        print(content_output)
        print("=" * 60)

        # Choose reranked docs (default: MMR, fallback: BM25, fallback: raw docs)
        chosen_docs = results[fastest_index].get("mmr_reranked") \
                      or results[fastest_index].get("bm25_reranked") \
                      or results[fastest_index]["documents"]

        # Build context for prompt
        context = "\n\n".join([doc.page_content for doc in chosen_docs])
        prompt_text = rag_prompt.format(context=context, question=query)

        # Get LLM response
        response = llm.invoke(prompt_text)
        print("\nLLM Response:\n", response.content)

        # Save to DOCX
        print("Entering the docx Module to save...")
        logger.info("Entering the docx Module to save...")
        save_to_docx(query, context, response.content, filename="artifacts/output.docx")

print("RAG PIPELINE EXECUTED SUCCESSFULLY")
logger.info("RAG PIPELINE EXECUTED SUCCESSFULLY")